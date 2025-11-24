"""Run management endpoints backed by the job registry."""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Request, status
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field, ConfigDict
from sqlalchemy.orm import Session

from server.core.modes import ResearchMode
from server.core.config import DATA_ROOT
from server.core.embeddings import ProfileEmbedder
from docx import Document

from ..services import JobRegistry, RunOptions
from ..dependencies import db_session, get_storage
from ..db import models
from ..storage.projects import ensure_project_structure
from ..adapters.storage import StorageBackend
from ..security import rbac
from .auth import SessionUser, require_roles


router = APIRouter(prefix="/deals/{deal_id}/runs", tags=["runs"])
oms_router = APIRouter(prefix="/deals/{deal_id}/oms", tags=["oms"])
run_router = APIRouter(prefix="/runs", tags=["runs"])


class CreateRunRequest(BaseModel):
    run_mode: str = Field("full", pattern="^(full|fast)$")
    research_mode: ResearchMode = ResearchMode.NONE
    interactive: bool = False
    deal_identifier: Optional[str] = None
    instructions: Optional[str] = None
    enable_vector_store: bool = True
    enable_web_search: bool = False
    included_file_ids: List[UUID] = Field(default_factory=list)
    parent_run_id: Optional[UUID] = None
    what_to_change: Optional[str] = None


class RunStatusResponse(BaseModel):
    id: UUID
    deal_id: str
    status: str
    run_mode: str
    research_mode: str
    created_at: datetime
    started_at: Optional[datetime] = None
    finished_at: Optional[datetime] = None
    result_path: Optional[str] = None
    error: Optional[str] = None
    params: dict
    instructions: Optional[str] = None
    included_file_ids: List[str] = Field(default_factory=list)
    parent_run_id: Optional[str] = None
    extracted_variables_artifact_id: Optional[UUID] = None


class DealOmResponse(BaseModel):
    run_id: UUID
    status: str
    created_at: datetime
    finished_at: Optional[datetime] = None
    rendered_artifact_id: Optional[UUID] = None
    variables_artifact_id: Optional[UUID] = None


class RunStepResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: UUID
    run_id: UUID
    name: str
    status: str
    started_at: Optional[datetime] = None
    finished_at: Optional[datetime] = None
    logs: Optional[str] = None


def _registry(request: Request) -> JobRegistry:
    registry = getattr(request.app.state, "job_registry", None)
    if registry is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Job registry unavailable")
    return registry


def _storage_key(deal_id: str, relative_path: str) -> str:
    clean = relative_path.replace("\\", "/").lstrip("/")
    return f"projects/{deal_id}/{clean}"


async def _ensure_artifact_local(
    deal_id: str,
    artifact: models.Artifact,
    storage: StorageBackend,
) -> Path:
    paths = ensure_project_structure(DATA_ROOT, deal_id)
    local_path = (paths.root / artifact.path).resolve()
    local_path.parent.mkdir(parents=True, exist_ok=True)

    if not local_path.exists():
        storage_key = _storage_key(deal_id, artifact.path)
        await run_in_threadpool(storage.download_to_path, storage_key, local_path)

    return local_path


def _markdown_to_docx_bytes(content: str) -> io.BytesIO:
    document = Document()
    in_code_block = False
    code_buffer: List[str] = []

    def _add_runs_with_emphasis(paragraph, text: str) -> None:
        # Minimal inline markdown support: **bold**, __bold__, `code`
        import re
        parts = re.split(r"(\*\*.+?\*\*|__.+?__|`.+?`)", text)
        for part in parts:
            if not part:
                continue
            if (part.startswith("**") and part.endswith("**") and len(part) >= 4):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif (part.startswith("__") and part.endswith("__") and len(part) >= 4):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif (part.startswith("`") and part.endswith("`") and len(part) >= 2):
                # Render inline code as a plain run for now
                paragraph.add_run(part[1:-1])
            else:
                paragraph.add_run(part)

    for raw_line in content.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                paragraph = document.add_paragraph("\n".join(code_buffer) if code_buffer else "")
                paragraph.style = "Intense Quote"
                code_buffer.clear()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            level = max(1, min(level, 4))
            heading_para = document.add_heading("", level=level)
            _add_runs_with_emphasis(heading_para, text or "")
            continue

        if stripped.startswith(('- ', '* ')):
            bullet_text = stripped[2:].strip()
            bullet_para = document.add_paragraph(style="List Bullet")
            _add_runs_with_emphasis(bullet_para, bullet_text)
            continue

        if stripped.startswith(">"):
            quote_para = document.add_paragraph("")
            quote_para.style = "Intense Quote"
            _add_runs_with_emphasis(quote_para, stripped[1:].strip())
            continue

        body_para = document.add_paragraph("")
        _add_runs_with_emphasis(body_para, stripped)

    if code_buffer:
        paragraph = document.add_paragraph("\n".join(code_buffer))
        paragraph.style = "Intense Quote"

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _get_rendered_artifact(db: Session, run_id: UUID) -> models.Artifact | None:
    return (
        db.query(models.Artifact)
        .filter(models.Artifact.run_id == run_id, models.Artifact.kind == "rendered_doc")
        .order_by(models.Artifact.created_at.desc())
        .first()
    )


def _get_variables_artifact(db: Session, run_id: UUID) -> models.Artifact | None:
    return (
        db.query(models.Artifact)
        .filter(models.Artifact.run_id == run_id, models.Artifact.kind == "variables")
        .order_by(models.Artifact.created_at.desc())
        .first()
    )


def _job_to_response(job: JobStatus) -> RunStatusResponse:
    data = job.to_dict()
    params = data.get("params", {}) or {}
    data["instructions"] = params.get("instructions_override")
    data["included_file_ids"] = params.get("included_file_ids", [])
    data["parent_run_id"] = params.get("parent_run_id")
    data["extracted_variables_artifact_id"] = params.get("extracted_variables_artifact_id")
    return RunStatusResponse(**data)


def _db_run_to_response(run: models.Run) -> RunStatusResponse:
    return RunStatusResponse(
        id=run.id,
        deal_id=str(run.deal_id),
        status=run.status,
        run_mode=run.mode,
        research_mode=run.research_mode,
        created_at=run.created_at,
        started_at=run.started_at,
        finished_at=run.finished_at,
        result_path=run.result_path,
        error=run.error,
        params=run.params or {},
        instructions=run.instructions,
        included_file_ids=run.included_file_ids or [],
        parent_run_id=str(run.parent_run_id) if run.parent_run_id else None,
        extracted_variables_artifact_id=run.extracted_variables_artifact_id,
    )


@router.get("/", response_model=List[RunStatusResponse])
async def list_runs(
    deal_id: str,
    request: Request,
    db: Session = Depends(db_session),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
) -> List[RunStatusResponse]:
    registry = _registry(request)
    jobs = registry.list_jobs(deal_id)
    job_map: dict[UUID, JobStatus] = {job.id: job for job in jobs}

    try:
        deal_uuid = UUID(deal_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Deal not found")

    rbac.ensure_deal_access(db, current_user, deal_uuid)

    responses: List[RunStatusResponse] = []

        db_runs = (
            db.query(models.Run)
        .filter(models.Run.deal_id == deal_uuid)
            .order_by(models.Run.created_at.desc())
            .all()
        )
        for run in db_runs:
            response = _db_run_to_response(run)
            job = job_map.pop(run.id, None)
            if job is not None:
                live = _job_to_response(job)
                response.status = live.status
                response.started_at = live.started_at or response.started_at
                response.finished_at = live.finished_at or response.finished_at
                response.result_path = live.result_path or response.result_path
                response.error = live.error or response.error
                response.params = live.params or response.params
                response.instructions = live.instructions or response.instructions
                response.included_file_ids = live.included_file_ids or response.included_file_ids
                response.parent_run_id = live.parent_run_id or response.parent_run_id
                response.extracted_variables_artifact_id = (
                    live.extracted_variables_artifact_id or response.extracted_variables_artifact_id
                )
            responses.append(response)

    # Include any remaining in-memory jobs (e.g., very new jobs not yet persisted)
    for job in job_map.values():
        responses.append(_job_to_response(job))

    # Sort by created_at desc for stable ordering
    responses.sort(key=lambda r: r.created_at, reverse=True)
    return responses


@router.post("/", response_model=RunStatusResponse, status_code=status.HTTP_201_CREATED)
async def create_run(
    deal_id: str,
    payload: CreateRunRequest,
    request: Request,
    db: Session = Depends(db_session),
    current_user: SessionUser = Depends(require_roles("editor", "admin")),
) -> RunStatusResponse:
    try:
        deal_uuid = UUID(deal_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Deal not found")
    rbac.ensure_deal_access(db, current_user, deal_uuid)
    registry = _registry(request)
    included_ids = [str(file_id) for file_id in payload.included_file_ids]
    parent_run_id = str(payload.parent_run_id) if payload.parent_run_id else None
    run_mode = payload.run_mode
    if parent_run_id:
        run_mode = "fast"
    options = RunOptions(
        interactive=payload.interactive,
        deal_identifier=payload.deal_identifier,
        run_mode=run_mode,
        research_mode=payload.research_mode.value,
        instructions_override=payload.instructions,
        enable_vector_store=payload.enable_vector_store,
        enable_web_search=payload.enable_web_search,
        included_file_ids=included_ids,
        parent_run_id=parent_run_id,
        variables_delta=payload.what_to_change,
    )
    job = registry.create_job(deal_id, options)
    data = job.to_dict()
    params = data.get("params", {}) or {}
    data["instructions"] = params.get("instructions_override")
    data["included_file_ids"] = params.get("included_file_ids", [])
    data["parent_run_id"] = params.get("parent_run_id")
    data["extracted_variables_artifact_id"] = params.get("extracted_variables_artifact_id")
    return RunStatusResponse(**data)


@router.get("/{run_id}", response_model=RunStatusResponse)
async def get_run(
    deal_id: str,
    run_id: UUID,
    request: Request,
    db: Session = Depends(db_session),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
) -> RunStatusResponse:
    registry = _registry(request)
    job = registry.get_job(run_id)
    if job is not None and job.deal_id == deal_id:
        rbac.ensure_deal_access(db, current_user, UUID(deal_id))
        return _job_to_response(job)

    run = db.get(models.Run, run_id)
    if run is not None and str(run.deal_id) == deal_id:
        rbac.ensure_deal_access(db, current_user, run.deal_id)
        return _db_run_to_response(run)

        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Run not found")


@oms_router.get("/", response_model=List[DealOmResponse])
async def list_deal_oms(
    deal_id: str,
    db: Session = Depends(db_session),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
) -> List[DealOmResponse]:
    try:
        deal_uuid = UUID(deal_id)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Deal not found")

    rbac.ensure_deal_access(db, current_user, deal_uuid)

    runs = (
        db.query(models.Run)
        .filter(models.Run.deal_id == deal_uuid)
        .order_by(models.Run.created_at.desc())
        .all()
    )
    run_ids = [run.id for run in runs]
    artifacts_map: dict[UUID, dict[str, Optional[UUID]]] = {
        run.id: {"rendered": None, "variables": None} for run in runs
    }

    if run_ids:
        artifacts = (
            db.query(models.Artifact)
            .filter(
                models.Artifact.run_id.in_(run_ids),
                models.Artifact.kind.in_(["rendered_doc", "variables"]),
            )
            .all()
        )
        for artifact in artifacts:
            slot = artifacts_map.get(artifact.run_id)
            if slot is None:
                continue
            if artifact.kind == "rendered_doc":
                slot["rendered"] = artifact.id
            elif artifact.kind == "variables":
                slot["variables"] = artifact.id

    responses: List[DealOmResponse] = []
    for run in runs:
        slots = artifacts_map.get(run.id) or {}
        responses.append(
            DealOmResponse(
                run_id=run.id,
                status=run.status,
                created_at=run.created_at,
                finished_at=run.finished_at,
                rendered_artifact_id=slots.get("rendered"),
                variables_artifact_id=slots.get("variables"),
            )
        )

    return responses


@oms_router.get("/{run_id}/download/{format}")
async def download_deal_om(
    deal_id: str,
    run_id: UUID,
    format: str,
    db: Session = Depends(db_session),
    storage: StorageBackend = Depends(get_storage),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
):
    run = db.get(models.Run, run_id)
    if not run or str(run.deal_id) != deal_id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Run not found for deal")
    if run.status != "success":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Run must be successful before exporting")

    rbac.ensure_deal_access(db, current_user, run.deal_id)

    artifact = _get_rendered_artifact(db, run_id)
    if artifact is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Rendered document not found")

    try:
        local_path = await _ensure_artifact_local(str(run.deal_id), artifact, storage)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to download artifact: {exc}")

    if not local_path.exists():
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Rendered document is unavailable")

    format_normalized = format.lower()
    if format_normalized == "md":
        filename = Path(artifact.path).name
        return FileResponse(
            str(local_path),
            filename=filename,
            media_type="text/markdown",
        )
    if format_normalized == "docx":
        try:
            content = local_path.read_text(encoding="utf-8")
        except Exception as exc:
            raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to read artifact: {exc}")

        buffer = _markdown_to_docx_bytes(content)
        filename_stem = Path(artifact.path).stem or f"run-{run.id}"
        docx_filename = f"{filename_stem}.docx"
        headers = {"Content-Disposition": f'attachment; filename="{docx_filename}"'}
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return StreamingResponse(buffer, media_type=media_type, headers=headers)

    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported format; use 'md' or 'docx'")


@run_router.get("/{run_id}", response_model=RunStatusResponse)
async def get_run_by_id(
    run_id: UUID,
    db: Session = Depends(db_session),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
) -> RunStatusResponse:
    run = rbac.ensure_run_access(db, current_user, run_id)
    return RunStatusResponse(
        id=run.id,
        deal_id=str(run.deal_id),
        status=run.status,
        run_mode=run.mode,
        research_mode=run.research_mode,
        created_at=run.created_at,
        started_at=run.started_at,
        finished_at=run.finished_at,
        result_path=run.result_path,
        error=run.error,
        params=run.params,
        instructions=run.instructions,
        included_file_ids=run.included_file_ids or [],
        parent_run_id=str(run.parent_run_id) if run.parent_run_id else None,
        extracted_variables_artifact_id=run.extracted_variables_artifact_id,
    )


@run_router.get("/{run_id}/steps", response_model=List[RunStepResponse])
async def get_run_steps(
    run_id: UUID,
    db: Session = Depends(db_session),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
) -> List[RunStepResponse]:
    run = rbac.ensure_run_access(db, current_user, run_id)
    steps = (
        db.query(models.RunStep)
        .filter(models.RunStep.run_id == run_id)
        .order_by(models.RunStep.started_at.asc())
        .all()
    )
    return [RunStepResponse.model_validate(step) for step in steps]


@run_router.get("/{run_id}/events")
async def stream_run_events(run_id: UUID) -> None:
    raise HTTPException(status_code=status.HTTP_501_NOT_IMPLEMENTED, detail="Run event stream not implemented yet")


@run_router.post("/{run_id}/embed", status_code=status.HTTP_201_CREATED)
async def embed_run_output(
    run_id: UUID,
    request: Request,
    db: Session = Depends(db_session),
    storage: StorageBackend = Depends(get_storage),
    current_user: SessionUser = Depends(require_roles("editor", "admin")),
):
    """
    Create a compact profile embedding for a run using its extracted variables.
    This mirrors the import system for consistent similarity matching.
    """
    vector_store = getattr(request.app.state, "vector_store", None)
    if vector_store is None:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail="Vector store unavailable")

    run = rbac.ensure_run_access(db, current_user, run_id)
    if run.status != "success":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Run must be successful before embedding")

    # Find the extracted variables artifact
    variables_artifact = (
        db.query(models.Artifact)
        .filter(models.Artifact.run_id == run_id, models.Artifact.kind == "variables")
        .order_by(models.Artifact.created_at.desc())
        .first()
    )

    if variables_artifact is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No extracted variables available for embedding")

    deal_uuid = None
    deal_id_str = str(run.deal_id)
    try:
        deal_uuid = UUID(deal_id_str)
    except Exception:
        deal_uuid = None

    # Load the extracted variables
    try:
        local_path = await _ensure_artifact_local(deal_id_str, variables_artifact, storage)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to download variables: {exc}")

    if not local_path.exists() or not local_path.is_file():
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Variables artifact is unavailable")

    try:
        import json
        with open(local_path, 'r', encoding='utf-8') as f:
            variables = json.load(f)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to parse variables: {exc}")

    # Build compact profile text directly from variables
    try:
        profile_text_parts = []
        title = variables.get("project_name") or variables.get("deal_name")
        if title:
            profile_text_parts.append(f"title:{str(title).strip()}")
        industry = variables.get("industry")
        if industry:
            profile_text_parts.append(f"industry:{str(industry).strip()}")
        project_type = variables.get("project_type") or variables.get("solution_type")
        if project_type:
            profile_text_parts.append(f"project_type:{str(project_type).strip()}")
        services = variables.get("services") or []
        if services:
            if isinstance(services, list):
                profile_text_parts.append("services:" + ",".join(str(x).strip() for x in services if x))
            else:
                profile_text_parts.append("services:" + str(services).strip())
        tags = variables.get("tags") or []
        if tags:
            if isinstance(tags, list):
                profile_text_parts.append("tags:" + ",".join(str(x).strip() for x in tags if x))
            else:
                profile_text_parts.append("tags:" + str(tags).strip())
        profile_text = " | ".join(profile_text_parts)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to build profile: {exc}")

    # Generate embedding from compact profile
    try:
        embedder = ProfileEmbedder()
        embedding = list(embedder.embed(profile_text))
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Embedding generation failed: {exc}")

    # Store full variables in metadata for reference
    metadata = {
        "deal_id": deal_id_str,
        "run_id": str(run.id),
        "mode": run.mode,
        "research_mode": run.research_mode,
        "result_path": run.result_path,
        "profile_text": profile_text,
        "title": variables.get("project_name"),
        "hours_total": variables.get("hours_total"),
        "timeline_weeks": variables.get("timeline_weeks"),
        "milestone_count": len(variables.get("milestones", [])),
        "services": variables.get("services"),
        "tags": variables.get("tags"),
        "dev_hours": variables.get("dev_hours"),
        "training_hours": variables.get("training_hours"),
        "pm_hours": variables.get("pm_hours"),
        "total_setup_cost": variables.get("total_setup_cost"),
        "monthly_operating_cost": variables.get("monthly_operating_cost"),
        "automation_outputs": variables.get("automation_outputs"),
        "client_name": variables.get("client_name"),
        "project_name": variables.get("project_name"),
        "industry": variables.get("industry"),
        "project_type": variables.get("project_type"),
    }
    if run.instructions:
        metadata["instructions"] = run.instructions

    try:
        embedding_id = vector_store.upsert_embedding(
            embedding=embedding,
        deal_id=deal_uuid,
            doc_kind="rendered_scope",
            metadata=metadata,
        )
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to store embedding: {exc}")

    return {"embedding_id": str(embedding_id), "profile_text": profile_text}


@run_router.get("/{run_id}/download-docx")
async def download_run_docx(
    run_id: UUID,
    db: Session = Depends(db_session),
    storage: StorageBackend = Depends(get_storage),
    current_user: SessionUser = Depends(require_roles("viewer", "editor", "admin")),
):
    run = rbac.ensure_run_access(db, current_user, run_id)
    if run.status != "success":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Run must be successful before exporting")

    artifact = _get_rendered_artifact(db, run_id)
    if artifact is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Rendered document not found")

    deal_id_str = str(run.deal_id)

    try:
        local_path = await _ensure_artifact_local(deal_id_str, artifact, storage)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to download artifact: {exc}")

    if not local_path.exists() or not local_path.is_file():
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Rendered document is unavailable")

    try:
        content = local_path.read_text(encoding="utf-8")
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Failed to read artifact: {exc}")

    buffer = _markdown_to_docx_bytes(content)
    filename_stem = Path(artifact.path).stem or f"run-{run.id}"
    docx_filename = f"{filename_stem}.docx"

    headers = {
        "Content-Disposition": f'attachment; filename="{docx_filename}"'
    }
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return StreamingResponse(buffer, media_type=media_type, headers=headers)

